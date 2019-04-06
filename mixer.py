import sys
from PyQt5.QtWidgets import (
    QApplication, QWidget, QGridLayout, QLabel,
    QLineEdit, QTextEdit, QDesktopWidget, QPushButton, QVBoxLayout
)
from docx import Document
from docx.shared import Inches
from PyQt5.QtGui import QIcon, QPixmap

class MyApp(QWidget):

    def __init__(self):
        super().__init__()
        self.files = []

        self.initUI()

    def initUI(self):

        # Center Align
        self.setWindowTitle('Centering')
        self.resize(500, 350)
        self.center()
        self.show()

        # Button
        btn1 = QPushButton('&실행', self)
        # btn1.setCheckable(True)
        # btn1.toggle()

        # Layout
        grid = QGridLayout()
        self.setLayout(grid)
        grid.addWidget(btn1, 0, 0)

        grid.addWidget(QLabel('글1:'), 1, 0)
        grid.addWidget(QLabel('글2:'), 2, 0)
        grid.addWidget(QLabel('글3:'), 3, 0)
        grid.addWidget(QLabel('글4:'), 4, 0)
        grid.addWidget(QLabel('글5:'), 5, 0)

        self.texts = [QLineEdit() for i in range(5)]

        grid.addWidget(self.texts[0], 1, 1)
        grid.addWidget(self.texts[1], 2, 1)
        grid.addWidget(self.texts[2], 3, 1)
        grid.addWidget(self.texts[3], 4, 1)
        grid.addWidget(self.texts[4], 5, 1)

        # 파일 랜덤 선택
        btn1.clicked.connect(self.proceed_random_mix)

        # 기본 로직은 이전에
        self.setWindowTitle('QGridLayout')
        self.center()
        # self.setGeometry(300, 300, 300, 200)
        self.show()

    def center(self):

        qr = self.frameGeometry()
        cp = QDesktopWidget().availableGeometry().center()
        qr.moveCenter(cp)
        self.move(qr.topLeft())

    def proceed_random_mix(self):
        self.choose_random_file()
        self.mix_and_extract_to_doc()

    def choose_random_file(self):
        import os, random
        for i in range(5,0,-1):
            file = random.choice(os.listdir(f"./images{i}/")) #change dir name to whatever
            self.files.append(file)

    def mix_and_extract_to_doc(self):
        document = Document()
        for i in range(1, 6):
            document.add_paragraph(self.texts[i-1].text())
            document.add_picture(f'./images{i}/'+self.files.pop(), width=Inches(4))

        document.save('랜덤 이미지 문서.docx')


if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = MyApp()
    sys.exit(app.exec_())
